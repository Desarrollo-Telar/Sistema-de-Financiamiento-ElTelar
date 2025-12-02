import os
from project.settings import MEDIA_ROOT, STATIC_ROOT, MEDIA_URL, STATIC_URL
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
from django.contrib.staticfiles import finders

from django.shortcuts import render, get_object_or_404, redirect

# MODELOS
from apps.financings.models import Recibo, Invoice,AccountStatement,Credit,PaymentPlan, Payment, Cuota
from apps.accountings.models import Creditor, Insurance
from apps.InvestmentPlan.models import InvestmentPlan
from apps.subsidiaries.models import Subsidiary


from django.conf import settings

from django.utils._os import safe_join
from django.core.exceptions import SuspiciousFileOperation
from weasyprint import HTML
# tiempo
from datetime import datetime,timedelta
from django.db.models import Q


from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .plan_pagos  import generar_estado_cuenta_word


def set_paragraph_format(p):
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0


def fecha_formateada(fecha):
    meses = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre",
    }
    
    return f"{fecha.day} de {meses[fecha.month]} del {fecha.year}"


def render_pagare_docx(request, id, customer_code):

    destino = get_object_or_404(InvestmentPlan, id=id)

    sucursal = destino.sucursal
    tasa_interes = destino.get_tasa() 
    plazo = destino.plazo if destino.plazo else 0
    cuota = destino.initial_amount

    if sucursal is None:
        sucursal = Subsidiary.objects.get(id=request.session['sucursal_id'])

    dia = datetime.now().date()

    cliente = destino.customer_id  # Ajusta esto si tu relación es distinta

    # ============================
    #   CREAR DOCUMENTO
    # ============================
    doc = Document()

    # Fuente general
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)

    # ============================
    #   ENCABEZADO CON LOGO
    # ============================
    logo_path = finders.find('img/el_telar.jpeg')

    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells

    # Columna izquierda → Logo
    if logo_path:
        row[0].paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.8))

    # Columna derecha → Encabezado centrado
    p = row[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("INVERSIONES INTEGRALES EL TELAR S.A.\n")
    r.bold = True
    r = p.add_run(f"PAGARÉ {destino.investment_plan_code}")
    r.bold = True

   
    

    # ============================
    #   TEXTO PRINCIPAL
    # ============================
    texto = (
        f"En la ciudad de {sucursal.get_direc().state}, departamento de {sucursal.get_direc().city} "
        f"el día {fecha_formateada(dia)}, YO, {cliente.get_full_name().upper()} de {cliente.get_edad_en_letras()} "
        f"({cliente.get_edad()}) años de edad, estado civil {cliente.get_estado_civil().upper()}, "
        f"de profesión {cliente.profession_trade.upper()}, {cliente.get_nacionalidad()}, de este domicilio me identifico "
        f"con el Documento de Identificación con Código Único de Identificación {cliente.get_numero_identificacion_en_letras()} "
        f"({cliente.formato_identificicaion()}) extendido por el Registro Nacional de las Personas. Por medio del presente "
        f"título de crédito consistente en un PAGARÉ, me comprometo de forma incondicional y sin necesidad "
        f"de cobro o requerimiento alguno pagar a Inversiones Integrales El Telar S.A. Únicamente mediante el depósito "
        f"correspondiente a la cuenta {sucursal.numero_de_cuenta_banco} del {sucursal.nombre_banco}, la cantidad "
        f"de {destino.en_letras_el_valor()} exactos (Q {destino.f_total_value_of_the_product_or_service()}), mediante {plazo} pagos "
        f"mensuales según la tabla que me es entregada adjunta a este documento."
    )
    

    p_texto = doc.add_paragraph(texto)
    set_paragraph_format(p_texto)
    doc.add_paragraph("")
    

    # ============================
    #   LISTA NUMERADA
    # ============================
    lista = [
        f"Pagar {tasa_interes}% de interés mensual, contabilizados a partir de la entrega de los fondos. ",

        f"En caso de incumplimiento al pago establecido, deberá pagar el equivalente al {tasa_interes}% mensual a la tasa de interés pactada a partir del primer día de morosidad. ",

        "Renuncio al fuero de mi domicilio y me someto a los tribunales que el tenedor del pagaré elija. ",

        "Todos los gastos que directa o inderectamente ocasione esta negociación son por mi cuenta, incluyendo la cobranza judicial o extrajudicial. "
        "El incumplimiento de este contrato dará derecho al tenedor a exigir el pago del saldo adeudado, para lo cual puede utilizar el presente documento como título ejecutivo o a su elección Certificación Contable del saldo adecuado. ",

        "ESTE PAGARÉ SE EMITE LIBRE DE PROTESTO, LIBRE DE FORMALIDADES DE PRESTACIÓN Y COBRO O REQUERIMIENTO. ",

        "En caso de juicio, ni el tenedor de este pagaré ni los auxiliares que proponga, estarán obligados a prestar garantía. ",

        "Acepto como buenas, líquidas y exigibles las cuentas que el tenedor del pagaré presente. ",

        
    ]


    for item in lista:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        

    doc.add_paragraph("")
    doc.add_paragraph("ACEPTO LIBRE DE PROTESTO:")
    
    # ============================
    #   DATOS FINALES DEL DEUDOR
    # ============================
    pf_1 = doc.add_paragraph(f"{cliente.get_full_name().upper()}")
    pf_1.bold = True
    set_paragraph_format(pf_1)
    pf_2 = doc.add_paragraph(f"DPI {cliente.identification_number}")
    set_paragraph_format(pf_2)
    pf_2 = doc.add_paragraph(f"{cliente.get_direccion().street.upper() if cliente.get_direccion() != '' else '' }")
    set_paragraph_format(pf_2)
    pf_4 = doc.add_paragraph(f"{sucursal.get_direc().state.upper()}, {sucursal.get_direc().city.upper()}")
    set_paragraph_format(pf_4)
    doc.add_paragraph("")
    doc.add_paragraph("__________________________   ÚLTIMA LINEA")
    doc.add_paragraph("")
    
    # ============================
    #   RESPUESTA HTTP
    # ============================
    generar_estado_cuenta_word(doc, id)
    
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="Pagare_{destino.investment_plan_code}.docx"'
    doc.save(response)

    
    return response
