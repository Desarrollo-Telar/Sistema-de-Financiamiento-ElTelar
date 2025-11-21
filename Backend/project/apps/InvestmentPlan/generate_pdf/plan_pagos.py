from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from datetime import datetime
from django.templatetags.static import static
import os
from django.contrib.staticfiles import finders

# Modelos
from apps.InvestmentPlan.models import InvestmentPlan
from apps.subsidiaries.models import Subsidiary

# Generando Plan de pagos
from apps.financings.clases.paymentplan import PaymentPlan
from apps.financings.clases.credit import Credit

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from django.http import HttpResponse
from django.contrib.staticfiles import finders
from django.shortcuts import get_object_or_404
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_paragraph_format(p):
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0

def set_table_border(table):
    tbl = table._element

    # Obtener o crear tblPr
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    # Crear elemento de bordes
    tblBorders = OxmlElement('w:tblBorders')

    # Configurar tipos de bordes
    for border_type in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')     # grosor del borde
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # negro
        tblBorders.append(border)

    tblPr.append(tblBorders)



def generar_estado_cuenta_word(doc, id):
    plan = get_object_or_404(InvestmentPlan, id=id)
    cliente = plan.customer_id
    sucursal = plan.sucursal
    dia = datetime.now().date()

    plazo = plan.plazo if plan.plazo else 1
    tasa_interes = plan.get_tasa()
    forma_pago = plan.forma_de_pago if plan.forma_de_pago else 'NIVELADA'
    fecha_inicio = plan.fecha_inicio if plan.fecha_inicio else dia

    credito = Credit('', plan.total_value_of_the_product_or_service, plazo, tasa_interes, forma_pago, 'MENSUAL', fecha_inicio.strftime('%Y-%m-%d'), 'CONSUMO', cliente)
    plan_pago = PaymentPlan(credito)

    cuotas = plan_pago.generar_plan()     # <--- AJUSTA si tu método es otro
    

    # ---------------------------
    # Crear documento
    # ---------------------------
    

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(10)

    # ---------------------------
    # ENCABEZADO
    # ---------------------------
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row = table.rows[0].cells

    # Logo
    logo_path = finders.find('img/el_telar.jpeg')
    if logo_path:
        row[0].paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.8))

    # Código del documento (derecha)
    p = row[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(plan.investment_plan_code).bold = True

    doc.add_paragraph("")  # pequeño espacio

    # ---------------------------
    # DATOS PRINCIPALES
    # ---------------------------
    

    table_info = doc.add_table(rows=4, cols=2)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    

    datos = [
        ("Deudor", f'{cliente.get_full_name().upper()}'),
        ("Monto Otorgado", f"Q{plan.total_value_of_the_product_or_service:,.2f}"),
        ("Fecha de Recibido", plan.fecha_inicio.strftime('%Y-%m-%d')),
        ("Forma de pago", f"{forma_pago}"),
    ]

    for i, (campo, valor) in enumerate(datos):
        c1 = table_info.rows[i].cells[0]
        c2 = table_info.rows[i].cells[1]

        c1.text = campo
        c2.text = valor


    

    doc.add_paragraph("")

    # ---------------------------
    # TABLA DE CUOTAS
    # ---------------------------
    tabla = doc.add_table(rows=1, cols=6)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tabla.rows[0].cells
    hdr[0].text = "No"
    hdr[1].text = "Fecha de Pago"
    hdr[2].text = "Saldo"
    hdr[3].text = "Intereses"
    hdr[4].text = "Capital"
    hdr[5].text = "Cuota"

    # Rellenar tabla
    for i, cuota in enumerate(cuotas, start=1):
        row = tabla.add_row().cells
        row[0].text = str(i)
        row[1].text = cuota['fecha_final'].strftime("%d/%m/%Y")
        row[2].text = f"Q {cuota['fmonto_prestado']}"
        row[3].text = f"Q {cuota['fintereses']}"
        row[4].text = f"Q {cuota['fcapital']}"
        row[5].text = f"Q {cuota['fcuota']}"
    
    
    set_table_border(tabla)
    # ---------------------------
    # TOTALES
    # ---------------------------
    doc.add_paragraph("")

    tabla_totales = doc.add_table(rows=4, cols=2)
    tot_rows = tabla_totales.rows

    tot = [
        ("Saldo Anterior", "Q0.00"),
        ("Gastos jurídicos", "Q0.00"),
        ("Otros Gastos", "Q0.00"),
        ("Líquido a Recibir", f"Q{plan.total_value_of_the_product_or_service:,.2f}"),
    ]

    for i, (campo, valor) in enumerate(tot):
        tot_rows[i].cells[0].text = campo
        tot_rows[i].cells[1].text = valor

    set_table_border(tabla_totales)

    # ---------------------------
    # DATOS DE DEPÓSITO
    # ---------------------------
    doc.add_paragraph("")

    tabla_dep = doc.add_table(rows=3, cols=2)
    dep = tabla_dep.rows

    dep[0].cells[0].text = "Depósito"
    dep[0].cells[1].text = sucursal.nombre_banco.upper()

    dep[1].cells[0].text = "Cuenta Monetaria"
    dep[1].cells[1].text = sucursal.numero_de_cuenta_banco

    dep[2].cells[0].text = "Nombre"
    dep[2].cells[1].text = "Inversiones Integrales el Telar S.A."
    set_table_border(tabla_dep)

    # Números de reporte:
    
    table_rep = doc.add_table(rows=1, cols=2)
    r = table_rep.rows[0].cells
    r[0].text = "Reportar pagos a los números:"
    r[1].text = f"{sucursal.numero_telefono}\n{sucursal.otro_numero_telefono}"
    set_table_border(table_rep)
    doc.add_paragraph("")
    
    doc.add_paragraph("FIRMA\n\n_____________________________")

    # ---------------------------
    # RESPUESTA HTTP
    # ---------------------------
    

